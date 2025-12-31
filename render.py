import io, random, datetime
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, PageBreak
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.pagesizes import A4
from reportlab.lib.colors import HexColor
from reportlab.lib.units import inch
from docx import Document

THEMES = ["#003366", "#7a1fa2", "#0b8457", "#b23a48"]

def clean_lines(text: str):
    # Preserve blank spacing while cleaning unwanted double spaces
    lines = []
    for l in text.split("\n"):
        if l.strip():
            lines.append(l.strip())
        else:
            lines.append("")   # keep blank line
    return lines

def header_footer(canvas, doc, rfq_id):
    canvas.saveState()
    canvas.setFont("Helvetica", 9)
    canvas.drawString(40, A4[1] - 30, rfq_id)
    canvas.drawRightString(A4[0] - 40, 30, f"Page {doc.page}")
    canvas.restoreState()

def render_pdf(rfq: dict, no: int) -> bytes:
    buf = io.BytesIO()
    rfq_id = f"RFQ_{no}_{rfq['name']}"
    doc = SimpleDocTemplate(buf, pagesize=A4,
        rightMargin=40,leftMargin=40,topMargin=60,bottomMargin=50)

    styles = getSampleStyleSheet()
    primary = HexColor(random.choice(THEMES))
    styles["Heading1"].textColor = primary
    styles["Heading2"].textColor = primary

    styles.add(ParagraphStyle(name="CoverTitle", fontSize=22,
        alignment=1, spaceAfter=30, textColor=primary))

    styles.add(ParagraphStyle(name="CoverMeta",
        fontSize=12, alignment=1, spaceAfter=10))

    styles.add(ParagraphStyle(name="Body", leading=14, spaceAfter=10))

    story = []
    issue_date = datetime.date.today().strftime("%d %B %Y")
    version = "v1.0"

    story.append(Spacer(1, 2 * inch))
    story.append(Paragraph("REQUEST FOR QUOTATION", styles["CoverTitle"]))
    story.append(Paragraph(rfq_id, styles["CoverMeta"]))
    story.append(Spacer(1, 30))
    story.append(Paragraph(f"<b>Domain:</b> {rfq.get('domain','N/A')}", styles["CoverMeta"]))
    story.append(Paragraph(f"<b>Category:</b> {rfq['name']}", styles["CoverMeta"]))
    story.append(Paragraph(f"<b>Document Version:</b> {version}", styles["CoverMeta"]))
    story.append(Paragraph(f"<b>Issue Date:</b> {issue_date}", styles["CoverMeta"]))
    story.append(PageBreak())

    story.append(Paragraph("TABLE OF CONTENTS", styles["Heading1"]))
    toc = ["Background & Objective","Scope of Work","Technical Requirements",
           "Service Level Agreement","Compliance & Standards",
           "Commercial Terms","Delivery Timeline","Evaluation Criteria",
           "Revision History"]
    for i,sec in enumerate(toc,1):
        story.append(Paragraph(f"{i}. {sec}", styles["Body"]))
    story.append(PageBreak())

    story.append(Paragraph(rfq_id, styles["Heading1"]))
    story.append(Spacer(1, 20))

    for line in clean_lines(rfq["body"]):
        if line.isupper() and len(line) < 60:
            story.append(Spacer(1, 14))
            story.append(Paragraph(line, styles["Heading2"]))
            story.append(Spacer(1, 6))
        else:
            story.append(Paragraph(line, styles["Body"]))

    story.append(PageBreak())
    story.append(Paragraph("REVISION HISTORY", styles["Heading1"]))
    story.append(Spacer(1, 10))
    story.append(Paragraph(f"v1.0 | {issue_date} | Initial Release", styles["Body"]))
    story.append(Paragraph("v1.1 | — | Reserved", styles["Body"]))
    story.append(Paragraph("v1.2 | — | Reserved", styles["Body"]))

    doc.build(story,
        onFirstPage=lambda c,d: header_footer(c,d,rfq_id),
        onLaterPages=lambda c,d: header_footer(c,d,rfq_id))

    buf.seek(0)
    return buf.read()

def render_docx(rfq: dict, no: int) -> bytes:
    rfq_id = f"RFQ_{no}_{rfq['name']}"
    issue_date = datetime.date.today().strftime("%d %B %Y")

    doc = Document()
    doc.add_heading("REQUEST FOR QUOTATION", level=0)
    doc.add_paragraph(rfq_id)
    doc.add_paragraph(f"Domain: {rfq.get('domain','N/A')}")
    doc.add_paragraph(f"Issue Date: {issue_date}")
    doc.add_page_break()

    doc.add_heading("TABLE OF CONTENTS", level=1)
    doc.add_paragraph("Update Table in Word")
    doc.add_page_break()

    doc.add_heading(rfq_id, level=1)
    for line in clean_lines(rfq["body"]):
        if line.isupper() and len(line) < 60:
            doc.add_heading(line, level=2)
        else:
            doc.add_paragraph(line)

    doc.add_page_break()
    doc.add_heading("REVISION HISTORY", level=1)
    doc.add_paragraph(f"v1.0 | {issue_date} | Initial Release")
    doc.add_paragraph("v1.1 | — | Reserved")
    doc.add_paragraph("v1.2 | — | Reserved")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
