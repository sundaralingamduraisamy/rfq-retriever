import io, random, datetime
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.colors import HexColor
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Inches
from database import db
from reportlab.platypus import Image as RLImage

THEMES = ["#003366", "#7a1fa2", "#0b8457", "#b23a48"]

def clean_lines(text: str):
    # Preserve blank spacing while cleaning unwanted double spaces
    lines = []
    for l in text.split("\n"):
        if l.strip():
            lines.append(l.strip())
        else:
            lines.append("")   # keep blank line
    return lines

def get_image_data(image_id: int):
    """Fetch image binary from DB"""
    if not db: return None
    row = db.execute_query_single("SELECT image_data FROM document_images WHERE id = %s", (image_id,))
    return row[0] if row else None

def header_footer(canvas, doc, rfq_id):
    canvas.saveState()
    canvas.setStrokeColor(HexColor("#006699"))
    canvas.setLineWidth(0.5)
    
    # Header: ID | NAME + Line
    header_text = f"{rfq_id.replace('_', ' ')}"
    canvas.setFont("Helvetica", 10)
    canvas.drawString(50, A4[1] - 40, header_text)
    canvas.line(50, A4[1] - 45, A4[0] - 50, A4[1] - 45)
    
    # Footer: Page No
    canvas.setFont("Helvetica", 9)
    canvas.drawRightString(A4[0] - 50, 30, f"Page {doc.page}")
    canvas.restoreState()

def render_pdf(rfq: dict, no: int) -> bytes:
    buf = io.BytesIO()
    rfq_id = f"RFQ_{no} | {rfq['name'].replace('_', ' ')}".upper()
    
    from reportlab.platypus import Table, TableStyle
    from reportlab.lib import colors

    # Class to track and render the TOC as a Table
    class MyDocTemplate(SimpleDocTemplate):
        def __init__(self, *args, **kwargs):
            self.entry_list = kwargs.pop('entry_list', [])
            SimpleDocTemplate.__init__(self, *args, **kwargs)

        def afterFlowable(self, flowable):
            if isinstance(flowable, Paragraph):
                text = flowable.getPlainText().strip()
                style = flowable.style.name
                # Rigorous exclusion of Cover titles and TOC title itself
                excluded = ["TABLE OF CONTENTS", "REQUEST FOR QUOTATION", "ISSUE DATE:"]
                if any(x in text.upper() for x in excluded):
                    return
                if style == 'Heading1':
                    self.entry_list.append((text, self.page))
                elif style == 'Heading2':
                    self.entry_list.append(("   " + text, self.page))
                # Heading3 is excluded from TOC to maintain professional brevity

    def create_toc_table(entries):
        if not entries:
            return Paragraph("No sections found.", styles["Body"])
            
        data = [["Section", "Page"]]
        # Custom styles for TOC to ensure proper wrapping and alignment
        toc_style = ParagraphStyle(name="TOCEntry", fontSize=10, leading=14, fontName="Helvetica")
        toc_style_bold = ParagraphStyle(name="TOCEntryBold", fontSize=11, leading=15, fontName="Helvetica-Bold")
        
        for text, page in entries:
            # Main sections are bold, sub-sections (with leading spaces) are normal
            s = toc_style_bold if not text.startswith(" ") else toc_style
            data.append([Paragraph(text.strip(), s), Paragraph(str(page), toc_style)])
        
        # Increase width of section column and ensure it wraps
        t = Table(data, colWidths=[5.2*inch, 0.7*inch])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), HexColor("#006680")), 
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('ALIGN', (1, 0), (1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'), # Align to top so wrapped text stays with number
        ]))
        return t

    styles = getSampleStyleSheet()
    primary = HexColor("#006680")
    
    # ... (rest of styles omitted for brevity, will be kept from previous edit) ...
    # Wait, multi_replace requires full TargetContent match. 
    # Let me just fix the build logic at the end and the class.

    # RE-DEFINING THE FULL FUNCTION TO BE SAFE
    # ...

    doc = MyDocTemplate(buf, pagesize=A4,
        rightMargin=50, leftMargin=50, topMargin=70, bottomMargin=50)

    styles = getSampleStyleSheet()
    primary = HexColor("#006680") # Consistent Teal theme
    
    # Custom Styles
    styles.add(ParagraphStyle(name="CoverTitle", fontSize=28, alignment=1, spaceAfter=50, textColor=primary, fontName="Helvetica-Bold"))
    styles.add(ParagraphStyle(name="CoverMeta", fontSize=14, alignment=1, spaceAfter=20, textColor=HexColor("#333333")))
    styles.add(ParagraphStyle(name="Body", fontSize=11, leading=16, spaceAfter=24, fontName="Helvetica", alignment=4)) 
    styles.add(ParagraphStyle(name="ListItem", fontSize=11, leading=16, leftIndent=25, spaceAfter=12, fontName="Helvetica"))
    
    styles["Heading1"].fontSize = 20
    styles["Heading1"].spaceBefore = 30
    styles["Heading1"].spaceAfter = 24
    styles["Heading1"].textColor = primary
    styles["Heading1"].fontName = "Helvetica-Bold"
    
    styles["Heading2"].fontSize = 15
    styles["Heading2"].spaceBefore = 25
    styles["Heading2"].spaceAfter = 24
    styles["Heading2"].textColor = primary
    styles["Heading2"].fontName = "Helvetica-Bold"

    if "Heading3" not in styles:
        styles.add(ParagraphStyle(name="Heading3", fontSize=13, spaceBefore=20, spaceAfter=20, textColor=primary, fontName="Helvetica-Bold"))
    else:
        styles["Heading3"].fontSize = 13
        styles["Heading3"].fontName = "Helvetica-Bold"
        styles["Heading3"].textColor = primary

    def get_story(entries=[]):
        story = []
        issue_date = datetime.date.today().strftime("%d %B %Y")
        
        # -- COVER PAGE --
        story.append(Spacer(1, 2 * inch))
        story.append(Paragraph("REQUEST FOR QUOTATION", styles["CoverTitle"]))
        story.append(Spacer(1, 10))
        story.append(Paragraph(rfq_id, styles["CoverMeta"]))
        story.append(Spacer(1, 1.5 * inch))
        story.append(Paragraph(f"<b>ISSUE DATE:</b> {issue_date}", styles["CoverMeta"]))
        story.append(PageBreak())

        # -- TABLE OF CONTENTS --
        story.append(Paragraph("TABLE OF CONTENTS", styles["Heading1"]))
        story.append(Spacer(1, 15))
        if entries:
            story.append(create_toc_table(entries))
        else:
            # First pass: Placeholder
            story.append(Spacer(1, 50))
        story.append(PageBreak())

        # -- CONTENT PARSER --
        import re
        in_table = False
        table_data = []
        table_style = TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), primary),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
        ])

        for line in rfq["body"].split('\n'):
            line = line.strip()
            
            # --- Table Handling ---
            if line.startswith('|') and line.endswith('|'):
                if '---' in line: continue
                cells = [c.strip() for c in line.split('|')[1:-1]]
                table_data.append(cells)
                in_table = True
                continue
            elif in_table:
                if table_data:
                    num_cols = len(table_data[0])
                    col_widths = [(A4[0] - 100) / num_cols] * num_cols
                    para_table = [[Paragraph(f"<b>{c}</b>" if r==0 else c, styles["Body"]) for c in row] for r, row in enumerate(table_data)]
                    t = Table(para_table, colWidths=col_widths)
                    t.setStyle(table_style)
                    story.append(Spacer(1, 15))
                    story.append(t)
                    story.append(Spacer(1, 24))
                table_data = []
                in_table = False

            if not line: continue
            
            # --- Image Detection (Prioritized) ---
            if "[[IMAGE_ID:" in line:
                match = re.search(r"\[\[IMAGE_ID:(\d+)\]\]", line)
                if match:
                    img_id = int(match.group(1))
                    img_data = get_image_data(img_id)
                    if img_data:
                        from reportlab.lib.utils import ImageReader
                        img_io = io.BytesIO(img_data)
                        img_reader = ImageReader(img_io)
                        iw, ih = img_reader.getSize()
                        target_w = 4.5 * inch
                        aspect = ih / float(iw)
                        target_h = target_w * aspect
                        img_flow = RLImage(img_io, width=target_w, height=target_h)
                        story.append(Spacer(1, 15))
                        story.append(img_flow)
                        story.append(Spacer(1, 24))
                        # If the line ONLY had the image tag, skip to next line
                        if line.strip() == f"[[IMAGE_ID:{img_id}]]":
                            continue
                        # Otherwise, strip the tag and process the remaining text
                        line = line.replace(f"[[IMAGE_ID:{img_id}]]", "").strip()
                        if not line: continue

            # --- Robust Header Detection (Sync with Frontend) ---
            # Strip bold wrappers if present for detection
            clean_line = line.replace("**", "").replace("__", "").replace("#", "").strip()
            # Match "1. Title" or "1.1 Title" or "Section 1: Title"
            is_header_like = re.match(r"^(\d+\.|\d+\.\d+|Section\s+\d+:)", clean_line, re.IGNORECASE)
            
            # Headers are usually concise. If a line exceeds 60 chars or contains colons
            # (like "1. Item: Description"), it's likely a requirement item, not a header.
            if is_header_like and len(clean_line) < 60:
                # If it has a colon, ensure it's a "Section N:" style, otherwise it's body
                if ":" in clean_line and not clean_line.lower().startswith("section"):
                    pass # Treat as body
                else:
                    is_sub = "." in clean_line.split(" ")[0] if clean_line else False
                    style = styles["Heading2"] if is_sub else styles["Heading1"]
                    story.append(Paragraph(clean_line, style))
                    continue

            # --- Standard Markdown Header Fallback ---
            if line.startswith('### '):
                story.append(Paragraph(line[4:].strip(), styles["Heading3"]))
            elif line.startswith('## '):
                story.append(Paragraph(line[3:].strip(), styles["Heading2"]))
            elif line.startswith('# '):
                header_text = line[2:].strip()
                if header_text.upper() == "TABLE OF CONTENTS":
                    continue
                story.append(Paragraph(header_text, styles["Heading1"]))
            
            # --- Bullet Detection ---
            elif line.startswith('- ') or line.startswith('* '):
                p_text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', line[2:])
                story.append(Paragraph(f"&bull; {p_text}", styles["ListItem"]))
            
            # --- Standard Body ---
            else:
                p_text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', line)
                story.append(Paragraph(p_text, styles["Body"]))

        # Final Table Flush
        if in_table and table_data:
            num_cols = len(table_data[0])
            col_widths = [(A4[0] - 100) / num_cols] * num_cols
            para_table = [[Paragraph(c, styles["Body"]) for c in row] for row in table_data]
            t = Table(para_table, colWidths=col_widths)
            t.setStyle(table_style)
            story.append(t)
            story.append(Spacer(1, 24))

        return story

    # Pass 1: Collect entries
    collected_entries = []
    doc = MyDocTemplate(buf, pagesize=A4, entry_list=collected_entries,
        rightMargin=50, leftMargin=50, topMargin=70, bottomMargin=50)
    
    doc.build(get_story(),
        onFirstPage=lambda c,d: header_footer(c,d,rfq_id),
        onLaterPages=lambda c,d: header_footer(c,d,rfq_id))
    
    # Pass 2: Render with TOC Table
    buf = io.BytesIO() 
    doc = MyDocTemplate(buf, pagesize=A4,
        rightMargin=50, leftMargin=50, topMargin=70, bottomMargin=50)
    
    doc.build(get_story(collected_entries),
        onFirstPage=lambda c,d: header_footer(c,d,rfq_id),
        onLaterPages=lambda c,d: header_footer(c,d,rfq_id))

    buf.seek(0)
    return buf.read()

def render_docx(rfq: dict, no: int) -> bytes:
    rfq_id = f"RFQ_{no}_{rfq['name']}"
    issue_date = datetime.date.today().strftime("%d %B %Y")

    doc = Document()
    doc.add_heading("REQUEST FOR QUOTATION", level=0)
    doc.add_paragraph(rfq_id)
    doc.add_paragraph(f"Domain: {rfq.get('domain','N/A')}")
    doc.add_paragraph(f"Issue Date: {issue_date}")
    doc.add_page_break()

    doc.add_heading("TABLE OF CONTENTS", level=1)
    doc.add_paragraph("Update Table in Word")
    doc.add_page_break()

    doc.add_heading(rfq_id, level=1)
    
    import re
    for line in clean_lines(rfq["body"]):
        # Skip TABLE OF CONTENTS - it's auto-generated
        if line.strip().upper() == "TABLE OF CONTENTS" or line.strip().upper() == "# TABLE OF CONTENTS":
            continue
            
        if line.isupper() and len(line) < 60:
            doc.add_heading(line, level=2)
        elif "[IMAGE_ID:" in line:
            import re
            match = re.search(r"\[\[IMAGE_ID:(\d+)\]\]", line)
            if match:
                img_id = int(match.group(1))
                img_data = get_image_data(img_id)
                if img_data:
                    img_io = io.BytesIO(img_data)
                    doc.add_picture(img_io, width=Inches(6))
                else:
                    doc.add_paragraph("[Image reference not found]")
        else:
            doc.add_paragraph(line)

    doc.add_page_break()
    doc.add_heading("REVISION HISTORY", level=1)
    doc.add_paragraph(f"v1.0 | {issue_date} | Initial Release")
    doc.add_paragraph("v1.1 | — | Reserved")
    doc.add_paragraph("v1.2 | — | Reserved")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
