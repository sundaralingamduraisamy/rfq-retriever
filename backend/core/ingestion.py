import fitz  # PyMuPDF
from datetime import datetime
from sentence_transformers import SentenceTransformer
from database import db
from settings import settings
from core.llm_provider import get_llm
from langchain_core.messages import HumanMessage, SystemMessage

# Initialize models
# embedding_model = SentenceTransformer(settings.EMBEDDING_MODEL_NAME) # Lazy loaded now
llm = get_llm()

from core.embedding_model import get_embedding_model
from core.image_processor import image_processor

class EmbeddingIndexer:
    """
    Extract → Summarize → Embed → Store
    """

    def __init__(self):
        # DB Tables are now created in main.py on startup
        pass

    def summarize(self, text: str) -> str:
        """Summarize document using configured LLM"""
        try:
            # Use full text (relying on model context window)
            messages = [
                SystemMessage(
                    content="You are a professional technical automotive specialist. Your task is to provide a HIGH-DETAIL, EXHAUSTIVE summary of the document. "
                            "Focus on: 1) Every specific technical parameter, 2) All mentioned ISO/SAE/IATF standards, 3) Detailed manufacturing/material requirements, "
                            "4) Complex constraints and 5) Precise timeline/commercial terms. "
                            "Maintain a dense, technical tone. DO NOT be concise; provide as much detail as possible for RAG purposes."
                ),
                HumanMessage(
                    content=f"Provide an exhaustive technical summary of this RFQ document:\n\n{text}"
                )
            ]

            response = llm.invoke(messages)
            return response.content
        except Exception as e:
            print(f"⚠️ Summarization failed: {e}")
            return text[:1000] # Fallback

    def index_document(self, filename: str, file_content: bytes, category: str = "General") -> dict:
        """
        Main indexing function:
        1. Store PDF binary in DB with category
        2. Extract Text
        3. Summarize
        4. Embed
        5. Store Vectors
        """
        try:
            print(f"\n📄 Indexing: {filename} (Category: {category})")

            # 1. Store/Update Document Record with category
            db.execute_update(
                """
                INSERT INTO documents (filename, category, file_size, file_content, uploaded_at)
                VALUES (%s, %s, %s, %s, %s)
                ON CONFLICT (filename) DO UPDATE 
                SET category = EXCLUDED.category,
                    file_size = EXCLUDED.file_size,
                    file_content = EXCLUDED.file_content,
                    uploaded_at = EXCLUDED.uploaded_at
                """,
                (filename, category, len(file_content), file_content, datetime.now())
            )

            # Get Document ID
            doc_id_row = db.execute_query_single("SELECT id FROM documents WHERE filename = %s", (filename,))
            if not doc_id_row:
                return {"success": False, "error": "DB record not found"}
            doc_id = doc_id_row[0]

            # 2. Extract Text based on file type
            file_ext = filename.split('.')[-1].lower()
            
            if file_ext == 'pdf':
                # PDF extraction using PyMuPDF
                doc = fitz.open(stream=file_content, filetype="pdf")
                full_text = ""
                for page in doc:
                    full_text += page.get_text()
                doc.close()
                
            elif file_ext == 'docx':
                # DOCX extraction using python-docx
                import docx
                import io
                doc = docx.Document(io.BytesIO(file_content))
                full_text = ""
                for paragraph in doc.paragraphs:
                    full_text += paragraph.text + "\n"
                # Also extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            full_text += cell.text + " "
                    full_text += "\n"
            elif file_ext in ['md', 'txt']:
                # Plain text / Markdown
                full_text = file_content.decode('utf-8', errors='ignore')
            else:
                print(f"❌ Unsupported file type: {file_ext}")
                return {"success": False, "error": "Unsupported file type"}

            # --- New: Image Context Injection for Generated RFQs ---
            # If this is a generated RFQ (markdown), it might contain [[IMAGE_ID:n]] tags.
            # We want to fetch those image descriptions and append them to the summary context
            # so that searching for "fuel pump image" finds this RFQ.
            import re
            image_tags = re.findall(r"\[\[IMAGE_ID:(\d+)\]\]", full_text)
            image_context = ""
            
            if image_tags:
                print(f"   🖼️ Found {len(image_tags)} image references. Fetching context...")
                unique_ids = list(set(image_tags))
                # Fetch descriptions for these IDs
                # We need to query the DB for these specific IDs
                placeholders = ','.join(['%s'] * len(unique_ids))
                img_rows = db.execute_query(f"SELECT id, description FROM document_images WHERE id IN ({placeholders})", tuple(unique_ids))
                
                if img_rows:
                    image_context = "\n\n--- REFERENCED IMAGES CONTEXT ---\n"
                    for r in img_rows:
                        image_context += f"[Image ID {r[0]}]: {r[1]}\n"
                    
                    # Append strictly for summarization/embedding (not modifying original file content)
                    print(f"   ✅ Injected context for {len(img_rows)} images into summary input.")
            
            # ------------------------------------------

            # --- New: Image Extraction (Standard Uploads) ---
            print(f"   📸 Extracting and filtering images...")
            processed_images = image_processor.process_content(file_content, file_ext)
            
            auto_count = sum(1 for img in processed_images if img["is_automobile"])
            non_auto_count = len(processed_images) - auto_count
            
            if auto_count > 0:
                print(f"   ✅ Found {auto_count} automobile-related images. Saving to DB...")
                image_processor.save_images_to_db(doc_id, processed_images)
            
            image_stats = {
                "total": len(processed_images),
                "automobile": auto_count,
                "non_automobile": non_auto_count,
                "has_automobile": auto_count > 0,
                "has_non_automobile": non_auto_count > 0
            }
            # ------------------------------------------

            if not full_text.strip():
                print("❌ No text extracted")
                return {"success": False, "error": "No text content", "image_stats": image_stats}

            # 3. Summarize
            print("   Generating summary...")
            # Append image context to the text being summarized so the LLM knows about the visuals
            summary_input = full_text + image_context
            summary = self.summarize(summary_input)
            print(f"   ✅ Model Summary:\n{summary}\n")
            
            # 4. Embed
            print("   Generating embedding...")
            model = get_embedding_model()
            embedding = model.encode(summary)
            embedding_str = str(embedding.tolist())

            # 5. Store Summary & Embedding
            # Store Summary
            db.execute_update(
                """
                INSERT INTO document_summaries (document_id, summary_text, word_count)
                VALUES (%s, %s, %s)
                ON CONFLICT (document_id) DO UPDATE 
                SET summary_text = EXCLUDED.summary_text,
                    word_count = EXCLUDED.word_count
                """,
                (doc_id, summary, len(summary.split()))
            )

            # Get Summary ID
            summary_id_row = db.execute_query_single("SELECT id FROM document_summaries WHERE document_id = %s", (doc_id,))
            if not summary_id_row:
                return {"success": False, "error": "Summary record not found", "image_stats": image_stats}
            summary_id = summary_id_row[0]

            # Store Vector
            db.execute_update(
                """
                INSERT INTO summary_embeddings (summary_id, embedding)
                VALUES (%s, %s)
                ON CONFLICT (summary_id) DO UPDATE 
                SET embedding = EXCLUDED.embedding
                """,
                (summary_id, embedding_str)
            )

            print(f"✅ Successfully indexed {filename}")
            return {"success": True, "image_stats": image_stats}

        except Exception as e:
            print(f"❌ Indexing error: {e}")
            import traceback
            traceback.print_exc()
            return {"success": False, "error": str(e)}

# Global instance
indexer = EmbeddingIndexer()
