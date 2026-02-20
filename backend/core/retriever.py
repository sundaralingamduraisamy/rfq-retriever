from datetime import datetime
from sentence_transformers import SentenceTransformer
from database import db
from settings import settings
import time
import torch

from core.embedding_model import get_embedding_model

# Initialize model lazily in functions
# embedding_model = SentenceTransformer(settings.EMBEDDING_MODEL_NAME)

def hybrid_search(query: str):
    """
    Search for documents using Vector Similarity on Summaries
    """
    start_time = time.time()
    
    if not db:
        return []

    try:
        # 1. Encode Query
        model = get_embedding_model()
        query_vec = model.encode(query)
        embedding_str = str(query_vec.tolist())

        # 2. Search in DB (Cosine Distance)
        # We search document_summaries via summary_embeddings table
        # 1 - (embedding <=> query) = Cosine Similarity
        # BOOST: If query text matches filename, give it a boost!
        sql_query = """
            SELECT 
                d.filename,
                ds.summary_text,
                (1 - (se.embedding <=> %s::vector)) * (CASE WHEN d.filename ILIKE %s THEN 1.2 ELSE 1.0 END) as similarity,
                ds.id as summary_id
            FROM summary_embeddings se
            JOIN document_summaries ds ON se.summary_id = ds.id
            JOIN documents d ON ds.document_id = d.id
            ORDER BY similarity DESC
            LIMIT %s
        """
        
        # Prepare params: embedding, query_pattern_for_boost, limit
        query_pattern = f"%{query.replace(' ', '%')}%"
        results = db.execute_query(sql_query, (embedding_str, query_pattern, settings.RETRIEVER_TOP_K))
        
        # 3. Format Results
        formatted_results = []
        for r in results:
            filename = r[0]
            summary_text = r[1]
            similarity = float(r[2])
            summary_id = r[3]

            formatted_results.append({
                "source": {
                    "file": filename,
                    "chunk_id": summary_id # Map summary_id to "chunk_id" for compatibility
                },
                "text": summary_text, # The AI sees the summary
                "relevance": round(similarity * 100, 2)
            })

        # print(f"✅ Search found {len(formatted_results)} results in {(time.time() - start_time)*1000:.1f}ms")
        return formatted_results

    except Exception as e:
        return []

def get_full_rfq(filename: str) -> str:
    """
    Retrieve full text content from DB for a given filename.
    Supports both PDF and DOCX files.
    """
    if not db:
        return "Database not connection."

    try:
        # Fetch binary content
        row = db.execute_query_single("SELECT file_content FROM documents WHERE filename = %s", (filename,))
        if not row:
            return "Document not found in database."
        
        file_content = bytes(row[0])
        file_ext = filename.split('.')[-1].lower()
        
        # Extract Text based on file type
        if file_ext == 'pdf':
            import fitz
            doc = fitz.open(stream=file_content, filetype="pdf")
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            
        elif file_ext == 'docx':
            import docx
            import io
            doc = docx.Document(io.BytesIO(file_content))
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                text += "\n"
        elif file_ext in ['md', 'txt']:
            # Plain text / Markdown
            text = file_content.decode('utf-8', errors='ignore')
        else:
            return f"Unsupported file type: {file_ext}"
            
        return text.strip()

    except Exception as e:
        return "Error reading document."

def search_images(query: str, top_k: int = 3):
    """
    Search for relevant images using CLIP text-image similarity.
    """
    if not db:
        return []

    try:
        from core.image_processor import get_model
        model, processor, mod_type = get_model()
        
        # 1. Encode Text Query using CLIP
        inputs = processor(text=[query], return_tensors="pt", padding=True)
        with torch.no_grad():
            text_features = model.get_text_features(**inputs)
        
        query_vec = text_features[0].tolist()
        embedding_str = str(query_vec)

        # 2. Search in DB - We fetch more (30) but will filter down
        # BOOST: Split keywords to handle typos (e.g., "suel shade" hits "SUNSHADE" via "shade")
        keywords = [k.strip() for k in query.split() if len(k.strip()) > 3]
        if not keywords: keywords = [query] # Fallback
        
        # Build dynamic ILIKE clauses
        boost_clauses = " OR ".join([f"d.filename ILIKE %s" for _ in keywords])
        boost_params = [f"%{k}%" for k in keywords]

        sql_query = f"""
            SELECT 
                di.id,
                di.description,
                d.filename,
                (1 - (ie.embedding <=> %s::vector)) * (CASE WHEN {boost_clauses or 'FALSE'} THEN 1.3 ELSE 1.0 END) as similarity,
                di.image_data
            FROM image_embeddings ie
            JOIN document_images di ON ie.image_id = di.id
            JOIN documents d ON di.document_id = d.id
            ORDER BY similarity DESC
            LIMIT 30
        """
        
        # Combined params: embedding, all keyword patterns
        results = db.execute_query(sql_query, (embedding_str, *boost_params))
        
        # Pass 1: Identify the "Primary Document" (Discovery Mode)
        primary_filename = None
        if results:
            # The #1 result's file is our primary target if it's reasonably confident.
            top_val = float(results[0][3])
            if top_val >= 0.15:
                primary_filename = results[0][2]

        # Pass 2: Fetch ALL images for the Primary Document to ensure "Total Images" requirement
        primary_images = []
        if primary_filename:
            # Note: We query the DB specifically for all images in this file
            sql_all_primary = """
                SELECT di.id, di.description, d.filename, 1.0 as similarity, di.image_data
                FROM document_images di
                JOIN documents d ON di.document_id = d.id
                WHERE d.filename = %s
            """
            primary_results = db.execute_query(sql_all_primary, (primary_filename,))
            for pr in primary_results:
                primary_images.append({
                    "id": pr[0],
                    "description": pr[1],
                    "file": pr[2],
                    "relevance": 100.0, # Target document images are 100% relevant context
                    "data": pr[4]
                })

        # Pass 3: IF PRIMARY HAS IMAGES, RETURN ONLY THOSE (Per User Request)
        if primary_images:
            # print(f"✅ Exclusive Discovery: Using {len(primary_images)} images specifically from {primary_filename}")
            return primary_images

        # Pass 4: FALLBACK - If primary document is image-less, collect high confidence global images
        other_images = []
        for r in results:
            similarity = float(r[3])
            # Strict threshold for fallback noise control
            if similarity >= 0.45: 
                other_images.append({
                    "id": r[0],
                    "description": r[1],
                    "file": r[2],
                    "relevance": round(similarity * 100, 2),
                    "data": r[4]
                })
            
        return other_images
    except Exception as e:
        return []
